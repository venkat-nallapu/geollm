"""
client/document_processor.py
Handles PDF and DOCX text extraction, cleaning, and chunking.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from loguru import logger

from src.utils.config import cfg


@dataclass
class DocumentChunk:
    text: str
    page_num: int
    chunk_index: int
    source: str   # filename


@dataclass
class ExtractedDocument:
    filename: str
    full_text: str
    chunks: list[DocumentChunk]
    num_pages: int


def _clean_text(text: str) -> str:
    """Remove excessive whitespace and fix common PDF extraction artifacts."""
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)   # dehyphenate line breaks
    return text.strip()


def _chunk_text(
    text: str,
    filename: str,
    page_num: int,
    chunk_size: int,
    overlap: int,
) -> list[DocumentChunk]:
    """Split page text into overlapping chunks by character count."""
    chunks = []
    start = 0
    idx = 0
    while start < len(text):
        end = start + chunk_size
        chunk_text = text[start:end]
        if chunk_text.strip():
            chunks.append(DocumentChunk(
                text=chunk_text,
                page_num=page_num,
                chunk_index=idx,
                source=filename,
            ))
        start += chunk_size - overlap
        idx += 1
    return chunks


def extract_pdf(file_bytes: bytes, filename: str) -> ExtractedDocument:
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber not installed. Run: pip install pdfplumber")

    max_pages = cfg("document", "max_pages", default=60)
    chunk_size = cfg("embedding", "chunk_size", default=512)
    overlap = cfg("embedding", "chunk_overlap", default=64)

    all_chunks: list[DocumentChunk] = []
    full_text_parts = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        num_pages = min(len(pdf.pages), max_pages)
        logger.info(f"Extracting {num_pages} pages from {filename}")

        for page_num, page in enumerate(pdf.pages[:num_pages], start=1):
            raw = page.extract_text() or ""
            cleaned = _clean_text(raw)
            if not cleaned:
                continue
            full_text_parts.append(f"--- Page {page_num} ---\n{cleaned}")
            page_chunks = _chunk_text(cleaned, filename, page_num, chunk_size, overlap)
            all_chunks.extend(page_chunks)

    full_text = "\n\n".join(full_text_parts)
    logger.info(f"PDF extraction complete: {len(all_chunks)} chunks, {num_pages} pages")
    return ExtractedDocument(
        filename=filename,
        full_text=full_text,
        chunks=all_chunks,
        num_pages=num_pages,
    )


def extract_docx(file_bytes: bytes, filename: str) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    chunk_size = cfg("embedding", "chunk_size", default=512)
    overlap = cfg("embedding", "chunk_overlap", default=64)

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = _clean_text("\n\n".join(paragraphs))
    # Treat entire docx as single "page" for chunking
    chunks = _chunk_text(full_text, filename, page_num=1, chunk_size=chunk_size, overlap=overlap)

    logger.info(f"DOCX extraction complete: {len(chunks)} chunks from {filename}")
    return ExtractedDocument(
        filename=filename,
        full_text=full_text,
        chunks=chunks,
        num_pages=1,
    )


def extract_txt(file_bytes: bytes, filename: str) -> ExtractedDocument:
    chunk_size = cfg("embedding", "chunk_size", default=512)
    overlap = cfg("embedding", "chunk_overlap", default=64)

    text = _clean_text(file_bytes.decode("utf-8", errors="replace"))
    chunks = _chunk_text(text, filename, page_num=1, chunk_size=chunk_size, overlap=overlap)

    logger.info(f"TXT extraction complete: {len(chunks)} chunks from {filename}")
    return ExtractedDocument(
        filename=filename,
        full_text=text,
        chunks=chunks,
        num_pages=1,
    )


def extract_document(file_bytes: bytes, filename: str) -> ExtractedDocument:
    """Route to correct extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_pdf(file_bytes, filename)
    elif ext in ("docx", "doc"):
        return extract_docx(file_bytes, filename)
    elif ext == "txt":
        return extract_txt(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")
